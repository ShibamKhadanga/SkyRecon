"""
SkyRecon â€“ Report Generator
Generates real PDF and DOCX reports from actual detection/disaster data.
"""

import os
import logging
from datetime import datetime
from typing import Optional

from sqlalchemy.orm import Session
from sqlalchemy import text

from ..core.config import settings
from .area_calculator import compute_coverage_stats

logger = logging.getLogger(__name__)

# Severity colors (RGB tuples for reportlab, hex for docx)
SEVERITY_COLORS = {
    1: {"rgb": (34, 197, 94),   "hex": "22c55e", "label": "Level 1 â€“ Minor"},
    2: {"rgb": (132, 204, 22),  "hex": "84cc16", "label": "Level 2 â€“ Low"},
    3: {"rgb": (234, 179, 8),   "hex": "eab308", "label": "Level 3 â€“ Moderate"},
    4: {"rgb": (249, 115, 22),  "hex": "f97316", "label": "Level 4 â€“ High"},
    5: {"rgb": (239, 68, 68),   "hex": "ef4444", "label": "Level 5 â€“ Critical"},
}


def _fetch_analysis(db: Session, analysis_id: int) -> Optional[dict]:
    row = db.execute(
        text("SELECT * FROM analyses WHERE id = :id"),
        {"id": analysis_id}
    ).first()
    if not row:
        return None
    return dict(row._mapping)


def _fetch_detections(db: Session, analysis_id: int) -> list[dict]:
    rows = db.execute(
        text("""
            SELECT d.*, c.name as category_name, c.color as category_color
            FROM detections d
            LEFT JOIN categories c ON d.category_id = c.id
            WHERE d.analysis_id = :id
            ORDER BY d.confidence DESC
        """),
        {"id": analysis_id}
    ).fetchall()
    return [dict(r._mapping) for r in rows]


def _fetch_disaster_events(db: Session, analysis_id: int) -> list[dict]:
    rows = db.execute(
        text("SELECT * FROM disaster_events WHERE analysis_id = :id ORDER BY severity DESC"),
        {"id": analysis_id}
    ).fetchall()
    return [dict(r._mapping) for r in rows]


def _fetch_screenshots(db: Session, analysis_id: int, report_type: str) -> list[dict]:
    """Fetch screenshot paths with timestamps for embedding in report."""
    if report_type == "disaster":
        rows = db.execute(
            text("""
                SELECT screenshot_path, timestamp, disaster_type, severity, confidence
                FROM disaster_events
                WHERE analysis_id = :id AND screenshot_path IS NOT NULL
                ORDER BY severity DESC
            """),
            {"id": analysis_id}
        ).fetchall()
        return [dict(r._mapping) for r in rows]
    else:
        rows = db.execute(
            text("""
                SELECT DISTINCT screenshot_path, timestamp, label, confidence
                FROM detections
                WHERE analysis_id = :id AND screenshot_path IS NOT NULL
                ORDER BY confidence DESC
                LIMIT 8
            """),
            {"id": analysis_id}
        ).fetchall()
        return [dict(r._mapping) for r in rows]


def _resolve_screenshot_path(rel_path: str) -> Optional[str]:
    """Convert /screenshots/filename.jpg to absolute filesystem path."""
    if not rel_path:
        return None
    filename = os.path.basename(rel_path)
    abs_path = os.path.join(settings.SCREENSHOTS_DIR, filename)
    return abs_path if os.path.exists(abs_path) else None


def _category_summary(detections: list[dict]) -> dict:
    """Groups detections by category and counts them."""
    summary = {}
    for det in detections:
        cat = det.get("category_name", "Unknown")
        summary[cat] = summary.get(cat, 0) + 1
    return dict(sorted(summary.items(), key=lambda x: x[1], reverse=True))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PDF GENERATOR
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def generate_pdf(analysis_id: int, report_type: str, db: Session) -> str:
    """Generates a PDF report and returns the file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether, Image as RLImage
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    os.makedirs(settings.REPORTS_DIR, exist_ok=True)
    filename = f"SkyRecon_{report_type}_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join(settings.REPORTS_DIR, filename)

    analysis = _fetch_analysis(db, analysis_id)
    if not analysis:
        raise ValueError(f"Analysis {analysis_id} not found")

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    # Custom styles
    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                  fontSize=22, textColor=colors.HexColor("#22c55e"),
                                  spaceAfter=4)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"],
                                     fontSize=10, textColor=colors.HexColor("#94a3b8"),
                                     spaceAfter=12)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"],
                                    fontSize=13, textColor=colors.HexColor("#f0fdf4"),
                                    spaceBefore=14, spaceAfter=6,
                                    backColor=colors.HexColor("#111827"),
                                    leftIndent=6, rightIndent=6, borderPad=4)
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.HexColor("#cbd5e1"),
                                 spaceAfter=4, leading=14)
    label_style = ParagraphStyle("Label", parent=styles["Normal"],
                                  fontSize=8, textColor=colors.HexColor("#64748b"))

    story = []

    # â”€â”€ Header â”€â”€
    story.append(Paragraph("SKYRECON", title_style))
    story.append(Paragraph("AI Powered Drone Intelligence Platform", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#22c55e"), spaceAfter=12))

    report_title = "Mapping & Survey Report" if report_type == "mapping" else "Disaster Assessment Report"
    story.append(Paragraph(report_title, heading_style))
    story.append(Spacer(1, 0.3*cm))

    # â”€â”€ Project Info Table â”€â”€
    created_at = analysis.get("created_at")
    completed_at = analysis.get("completed_at")
    proc_time = analysis.get("processing_time")

    info_data = [
        ["Project Name", analysis.get("project_name", "â€”")],
        ["Location", analysis.get("location") or "â€”"],
        ["Drone Model", analysis.get("drone_model") or "â€”"],
        ["Detection Mode", analysis.get("detection_mode", "standard").upper()],
        ["Analysis Date", created_at.strftime("%d %b %Y, %H:%M") if created_at else "â€”"],
        ["Processing Time", f"{proc_time:.1f} seconds" if proc_time else "â€”"],
        ["Status", analysis.get("status", "â€”").upper()],
    ]
    if analysis.get("description"):
        info_data.append(["Description", analysis["description"]])

    info_table = Table(info_data, colWidths=[4*cm, 13*cm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#0c1017")),
        ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#06080d")),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#22c55e")),
        ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#e2e8f0")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e293b")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#0c1017"), colors.HexColor("#0a0e16")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.5*cm))

    # â”€â”€ MAPPING REPORT CONTENT â”€â”€
    if report_type == "mapping":
        detections = _fetch_detections(db, analysis_id)
        cat_summary = _category_summary(detections)
        coverage = compute_coverage_stats(detections)

        story.append(Paragraph("Detection Summary", heading_style))

        # Stats row
        stats_data = [
            ["Unique Objects", "Categories Found", "Area Covered", "Coverage %"],
            [
                str(len(detections)),
                str(len(cat_summary)),
                f"{coverage['total_area_sqm']} mÂ²",
                f"{coverage['coverage_percent']}%",
            ]
        ]
        stats_table = Table(stats_data, colWidths=[4.25*cm]*4)
        stats_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#0c1017")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#64748b")),
            ("TEXTCOLOR", (0, 1), (-1, 1), colors.HexColor("#22c55e")),
            ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e293b")),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.4*cm))

        # Category breakdown table
        if cat_summary:
            story.append(Paragraph("Category Breakdown", heading_style))
            cat_data = [["Category", "Count", "% of Total"]]
            total = len(detections)
            for cat, count in cat_summary.items():
                pct = f"{(count/total*100):.1f}%" if total > 0 else "0%"
                cat_data.append([cat, str(count), pct])

            cat_table = Table(cat_data, colWidths=[8*cm, 4*cm, 5*cm])
            cat_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#22c55e")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#0c1017"), colors.HexColor("#06080d")]),
                ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#cbd5e1")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e293b")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ]))
            story.append(cat_table)
            story.append(Spacer(1, 0.4*cm))

        # Top detections
        if detections:
            story.append(Paragraph("Top Detections (by Confidence)", heading_style))
            det_data = [["Label", "Category", "Confidence", "Timestamp"]]
            for det in detections[:20]:
                det_data.append([
                    det.get("label", "â€”")[:40],
                    det.get("category_name", "â€”"),
                    f"{det.get('confidence', 0):.0%}",
                    f"{det.get('timestamp', 0):.1f}s" if det.get("timestamp") else "â€”",
                ])
            det_table = Table(det_data, colWidths=[7*cm, 4*cm, 3*cm, 3*cm])
            det_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#22c55e")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#0c1017"), colors.HexColor("#06080d")]),
                ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#cbd5e1")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e293b")),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(det_table)
            story.append(Spacer(1, 0.4*cm))

        # AI Recommendations
        story.append(Paragraph("AI Planning Recommendations", heading_style))
        empty_pct = 100 - coverage["coverage_percent"]
        rec_text = (
            f"Analysis detected <b>{len(detections)}</b> objects across <b>{len(cat_summary)}</b> categories. "
            f"Total covered area: <b>{coverage['total_area_sqm']} mÂ²</b> "
            f"({coverage['coverage_percent']}% of surveyed zone). "
            f"Estimated available area: <b>{coverage['empty_area_sqm']} mÂ²</b> ({empty_pct:.1f}%). "
        )
        if "Plants" in cat_summary or "Trees" in cat_summary:
            plant_count = cat_summary.get("Plants", 0) + cat_summary.get("Trees", 0)
            rec_text += (
                f"Vegetation analysis: {plant_count} plant/tree instances detected. "
                f"Recommend planting in the {coverage['empty_area_sqm']} mÂ² of open soil "
                f"to improve green cover index."
            )
        story.append(Paragraph(rec_text, body_style))

        # â”€â”€ Screenshots Section â”€â”€
        screenshots = _fetch_screenshots(db, analysis_id, "mapping")
        if screenshots:
            story.append(Spacer(1, 0.4*cm))
            story.append(Paragraph("Detection Screenshots", heading_style))
            story.append(Paragraph(
                "Annotated frames captured during AI analysis. Timestamp shown on each image indicates exact position in the video.",
                body_style
            ))
            story.append(Spacer(1, 0.3*cm))
            # 2 screenshots per row
            img_pairs = [screenshots[i:i+2] for i in range(0, len(screenshots), 2)]
            for pair in img_pairs:
                row_items = []
                for shot in pair:
                    abs_path = _resolve_screenshot_path(shot.get("screenshot_path", ""))
                    if abs_path:
                        try:
                            img = RLImage(abs_path, width=8*cm, height=5*cm)
                            ts = shot.get('timestamp', 0)
                            label = shot.get('label', 'Detection')
                            conf = shot.get('confidence', 0)
                            caption = f"{label[:30]} | {conf:.0%} conf | T={ts:.1f}s"
                            row_items.append([img, Paragraph(caption, label_style)])
                        except Exception:
                            pass
                if row_items:
                    if len(row_items) == 2:
                        tbl = Table([[row_items[0][0], row_items[1][0]],
                                     [row_items[0][1], row_items[1][1]]],
                                    colWidths=[8.5*cm, 8.5*cm])
                    else:
                        tbl = Table([[row_items[0][0]], [row_items[0][1]]], colWidths=[17*cm])
                    tbl.setStyle(TableStyle([
                        ("ALIGN", (0,0), (-1,-1), "CENTER"),
                        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                        ("TOPPADDING", (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 0.3*cm))

    # â”€â”€ DISASTER REPORT CONTENT â”€â”€
    else:
        events = _fetch_disaster_events(db, analysis_id)

        story.append(Paragraph("Disaster Events Detected", heading_style))

        if not events:
            story.append(Paragraph("No confirmed disaster events detected in this footage.", body_style))
        else:
            for event in events:
                sev = event.get("severity", 1)
                sev_info = SEVERITY_COLORS.get(sev, SEVERITY_COLORS[1])
                sev_color = colors.HexColor(f"#{sev_info['hex']}")

                event_data = [
                    ["Disaster Type", event.get("disaster_type", "â€”").upper()],
                    ["Severity", sev_info["label"]],
                    ["Confidence", f"{event.get('confidence', 0):.0%}"],
                    ["Affected Area", f"{event.get('affected_area', 0):.0f} mÂ²"],
                    ["Timestamp", f"{event.get('timestamp', 0):.1f}s into video"],
                    ["Recommendations", event.get("recommendations", "â€”")],
                ]
                if event.get("resource_estimation"):
                    res = event["resource_estimation"]
                    event_data.append(["Resources Needed",
                        f"Rescue Teams: {res.get('rescue_teams',0)} | "
                        f"Ambulances: {res.get('ambulances',0)} | "
                        f"Boats: {res.get('rescue_boats',0)} | "
                        f"Staff: {res.get('support_staff',0)}"
                    ])

                ev_table = Table(event_data, colWidths=[4*cm, 13*cm])
                ev_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#0c1017")),
                    ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#06080d")),
                    ("TEXTCOLOR", (0, 0), (0, -1), sev_color),
                    ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#e2e8f0")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#1e293b")),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1),
                     [colors.HexColor("#0c1017"), colors.HexColor("#0a0e16")]),
                ]))

                # Embed screenshot for this event if available
                block = [ev_table, Spacer(1, 0.2*cm)]
                abs_path = _resolve_screenshot_path(event.get("screenshot_path", ""))
                if abs_path:
                    try:
                        img = RLImage(abs_path, width=12*cm, height=7*cm)
                        ts = event.get('timestamp', 0)
                        dtype = event.get('disaster_type', '').upper()
                        caption = Paragraph(
                            f"Screenshot: {dtype} detected at T={ts:.1f}s | "
                            f"Severity: {sev_info['label']} | "
                            f"Confidence: {event.get('confidence', 0):.0%}",
                            ParagraphStyle("Caption", parent=styles["Normal"],
                                           fontSize=8, textColor=sev_color,
                                           alignment=TA_CENTER)
                        )
                        block += [img, caption, Spacer(1, 0.2*cm)]
                    except Exception:
                        pass

                story.append(KeepTogether(block + [Spacer(1, 0.4*cm)]))

    # â”€â”€ Footer â”€â”€
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#1e293b")))
    story.append(Paragraph(
        f"Generated by SkyRecon AI Drone Intelligence Platform â€¢ {datetime.now().strftime('%d %b %Y %H:%M')}",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7,
                       textColor=colors.HexColor("#475569"), alignment=TA_CENTER)
    ))

    doc.build(story)
    logger.info(f"[REPORT] PDF generated: {filepath}")
    return f"/reports/{filename}"


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  DOCX GENERATOR
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def generate_docx(analysis_id: int, report_type: str, db: Session) -> str:
    """Generates a DOCX report and returns the file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs(settings.REPORTS_DIR, exist_ok=True)
    filename = f"SkyRecon_{report_type}_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(settings.REPORTS_DIR, filename)

    analysis = _fetch_analysis(db, analysis_id)
    if not analysis:
        raise ValueError(f"Analysis {analysis_id} not found")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_heading(text, level=1, color_hex="22c55e"):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16)
            )
        return p

    def add_para(text, bold=False, color_hex=None, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color_hex:
            run.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16)
            )
        return p

    def add_table_row(table, cells, bold_first=True):
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            row.cells[i].text = str(cell_text)
            if bold_first and i == 0:
                for para in row.cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True

    # â”€â”€ Title â”€â”€
    add_heading("SKYRECON", level=1, color_hex="22c55e")
    add_para("AI Powered Drone Intelligence Platform", color_hex="94a3b8")
    doc.add_paragraph("â”€" * 80)

    report_title = "Mapping & Survey Report" if report_type == "mapping" else "Disaster Assessment Report"
    add_heading(report_title, level=2, color_hex="f0fdf4")

    # â”€â”€ Project Info â”€â”€
    add_heading("Project Information", level=3, color_hex="22c55e")
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"

    created_at = analysis.get("created_at")
    proc_time = analysis.get("processing_time")

    info_rows = [
        ("Project Name", analysis.get("project_name", "â€”")),
        ("Location", analysis.get("location") or "â€”"),
        ("Drone Model", analysis.get("drone_model") or "â€”"),
        ("Detection Mode", analysis.get("detection_mode", "standard").upper()),
        ("Analysis Date", created_at.strftime("%d %b %Y, %H:%M") if created_at else "â€”"),
        ("Processing Time", f"{proc_time:.1f} seconds" if proc_time else "â€”"),
        ("Status", analysis.get("status", "â€”").upper()),
    ]
    if analysis.get("description"):
        info_rows.append(("Description", analysis["description"]))

    for label, value in info_rows:
        add_table_row(info_table, [label, value])

    doc.add_paragraph()

    if report_type == "mapping":
        detections = _fetch_detections(db, analysis_id)
        cat_summary = _category_summary(detections)
        coverage = compute_coverage_stats(detections)

        add_heading("Detection Summary", level=3, color_hex="22c55e")
        add_para(
            f"Unique Objects Tracked: {len(detections)}  |  "
            f"Categories: {len(cat_summary)}  |  "
            f"Area Covered: {coverage['total_area_sqm']} mÂ²  |  "
            f"Coverage: {coverage['coverage_percent']}%",
            bold=True
        )
        doc.add_paragraph()

        if cat_summary:
            add_heading("Category Breakdown", level=3, color_hex="22c55e")
            cat_table = doc.add_table(rows=1, cols=3)
            cat_table.style = "Table Grid"
            hdr = cat_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Category", "Count", "% of Total"
            total = len(detections)
            for cat, count in cat_summary.items():
                row = cat_table.add_row()
                row.cells[0].text = cat
                row.cells[1].text = str(count)
                row.cells[2].text = f"{(count/total*100):.1f}%" if total > 0 else "0%"
            doc.add_paragraph()

        if detections:
            add_heading("Top Detections (by Confidence)", level=3, color_hex="22c55e")
            det_table = doc.add_table(rows=1, cols=4)
            det_table.style = "Table Grid"
            hdr = det_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Label", "Category", "Confidence", "Timestamp"
            for det in detections[:20]:
                row = det_table.add_row()
                row.cells[0].text = (det.get("label") or "â€”")[:40]
                row.cells[1].text = det.get("category_name") or "â€”"
                row.cells[2].text = f"{det.get('confidence', 0):.0%}"
                row.cells[3].text = f"{det.get('timestamp', 0):.1f}s" if det.get("timestamp") else "â€”"
            doc.add_paragraph()

        add_heading("AI Planning Recommendations", level=3, color_hex="22c55e")
        empty_pct = 100 - coverage["coverage_percent"]
        rec = (
            f"Analysis detected {len(detections)} objects across {len(cat_summary)} categories. "
            f"Total covered area: {coverage['total_area_sqm']} mÂ² ({coverage['coverage_percent']}%). "
            f"Available area: {coverage['empty_area_sqm']} mÂ² ({empty_pct:.1f}%). "
        )
        if "Plants" in cat_summary or "Trees" in cat_summary:
            plant_count = cat_summary.get("Plants", 0) + cat_summary.get("Trees", 0)
            rec += (
                f"Vegetation: {plant_count} plant/tree instances detected. "
                f"Recommend planting in {coverage['empty_area_sqm']} mÂ² of open soil."
            )
        add_para(rec)

        # â”€â”€ Screenshots â”€â”€
        screenshots = _fetch_screenshots(db, analysis_id, "mapping")
        if screenshots:
            doc.add_paragraph()
            add_heading("Detection Screenshots", level=3, color_hex="22c55e")
            add_para("Annotated frames captured during AI analysis. Timestamp on each image indicates exact position in the video.")
            for shot in screenshots:
                abs_path = _resolve_screenshot_path(shot.get("screenshot_path", ""))
                if abs_path:
                    try:
                        ts = shot.get('timestamp', 0)
                        label = shot.get('label', 'Detection')
                        conf = shot.get('confidence', 0)
                        doc.add_picture(abs_path, width=Cm(14))
                        add_para(f"{label[:40]} | Confidence: {conf:.0%} | Timestamp: T={ts:.1f}s",
                                 color_hex="64748b")
                        doc.add_paragraph()
                    except Exception:
                        pass

    else:
        events = _fetch_disaster_events(db, analysis_id)
        add_heading("Disaster Events Detected", level=3, color_hex="ef4444")

        if not events:
            add_para("No confirmed disaster events detected in this footage.")
        else:
            for event in events:
                sev = event.get("severity", 1)
                sev_info = SEVERITY_COLORS.get(sev, SEVERITY_COLORS[1])
                hex_c = sev_info["hex"]

                add_heading(
                    f"{event.get('disaster_type','â€”').upper()} â€” {sev_info['label']}",
                    level=4, color_hex=hex_c
                )
                ev_table = doc.add_table(rows=0, cols=2)
                ev_table.style = "Table Grid"
                res = event.get("resource_estimation") or {}
                rows_data = [
                    ("Disaster Type", event.get("disaster_type", "â€”").upper()),
                    ("Severity", sev_info["label"]),
                    ("Confidence", f"{event.get('confidence', 0):.0%}"),
                    ("Affected Area", f"{event.get('affected_area', 0):.0f} mÂ²"),
                    ("Timestamp", f"{event.get('timestamp', 0):.1f}s into video"),
                    ("Rescue Teams", str(res.get("rescue_teams", 0))),
                    ("Ambulances", str(res.get("ambulances", 0))),
                    ("Rescue Boats", str(res.get("rescue_boats", 0))),
                    ("Support Staff", str(res.get("support_staff", 0))),
                    ("Recommendations", event.get("recommendations", "â€”")),
                ]
                for label, value in rows_data:
                    add_table_row(ev_table, [label, value])

                # Embed screenshot for this event
                abs_path = _resolve_screenshot_path(event.get("screenshot_path", ""))
                if abs_path:
                    try:
                        doc.add_paragraph()
                        doc.add_picture(abs_path, width=Cm(14))
                        ts = event.get('timestamp', 0)
                        dtype = event.get('disaster_type', '').upper()
                        add_para(
                            f"Screenshot: {dtype} at T={ts:.1f}s | "
                            f"{sev_info['label']} | "
                            f"Confidence: {event.get('confidence', 0):.0%}",
                            color_hex=hex_c
                        )
                    except Exception:
                        pass
                doc.add_paragraph()

    # Footer
    doc.add_paragraph("â”€" * 80)
    add_para(
        f"Generated by SkyRecon AI Drone Intelligence Platform â€¢ {datetime.now().strftime('%d %b %Y %H:%M')}",
        color_hex="475569"
    )

    doc.save(filepath)
    logger.info(f"[REPORT] DOCX generated: {filepath}")
    return f"/reports/{filename}"


def generate_report(analysis_id: int, report_type: str, fmt: str, db: Session) -> str:
    """Entry point â€” dispatches to PDF or DOCX generator."""
    if fmt == "pdf":
        return generate_pdf(analysis_id, report_type, db)
    elif fmt == "docx":
        return generate_docx(analysis_id, report_type, db)
    else:
        raise ValueError(f"Unsupported format: {fmt}")

